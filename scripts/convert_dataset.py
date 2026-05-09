#!/usr/bin/env python3
"""Convert .docx and .xlsx source data into markdown the LLM can consume.

Usage:

    # Convert one dataset's raw/ folder into ../{dataset}/*.md
    python scripts/convert_dataset.py --in backend/data/asurion/raw --out backend/data/asurion

    # Convert the standalone PRD template
    python scripts/convert_dataset.py --in backend/data/sprntly_prd_template.docx \\
        --out backend/data

Behavior:
- .docx → .md preserves headings (Heading 1 → '#'), paragraphs, and tables.
- .xlsx → .md emits one section per sheet, truncating sheets longer than 30 rows.
- The answer key (asurion_expected_output.*) is routed into a `_reference/`
  subfolder so it never gets fed back to the LLM accidentally.

Requires:
    pip install python-docx openpyxl

Run from the repo root.
"""
from __future__ import annotations

import argparse
from pathlib import Path

import docx
import openpyxl


def docx_to_md(path: Path) -> str:
    doc = docx.Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text
        if not text.strip():
            parts.append("")
            continue
        if style.startswith("Heading 1"):
            parts.append(f"# {text}")
        elif style.startswith("Heading 2"):
            parts.append(f"## {text}")
        elif style.startswith("Heading 3"):
            parts.append(f"### {text}")
        elif style.startswith("Heading"):
            parts.append(f"#### {text}")
        else:
            parts.append(text)
    for ti, table in enumerate(doc.tables):
        parts.append(f"\n[TABLE {ti}]")
        for row in table.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
    return "\n".join(parts)


def xlsx_to_md(path: Path, max_rows: int = 30) -> str:
    wb = openpyxl.load_workbook(str(path), data_only=True)
    out: list[str] = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        out.append(f"## Sheet: {sn}  (rows={ws.max_row}, cols={ws.max_column})")
        for ri, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if ri > max_rows:
                out.append(f"... [truncated, {ws.max_row - max_rows} more rows]")
                break
            cells = [str(c) if c is not None else "" for c in row]
            out.append("| " + " | ".join(cells) + " |")
        out.append("")
    return "\n".join(out)


REFERENCE_PATTERNS = ("expected_output",)


def is_reference(name: str) -> bool:
    return any(p in name for p in REFERENCE_PATTERNS)


def convert_one(src: Path, out_dir: Path) -> Path:
    suffix = src.suffix.lower()
    if suffix == ".docx":
        text = docx_to_md(src)
    elif suffix == ".xlsx":
        text = xlsx_to_md(src)
    else:
        raise ValueError(f"Unsupported extension: {suffix}")
    target_dir = out_dir / "_reference" if is_reference(src.stem) else out_dir
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / (src.stem + ".md")
    target.write_text(text)
    return target


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="src", required=True, help="File or directory to convert")
    ap.add_argument("--out", dest="dst", required=True, help="Output directory")
    args = ap.parse_args()

    src = Path(args.src).expanduser().resolve()
    dst = Path(args.dst).expanduser().resolve()
    if not src.exists():
        raise SystemExit(f"Input not found: {src}")
    dst.mkdir(parents=True, exist_ok=True)

    if src.is_file():
        out = convert_one(src, dst)
        print(f"Wrote {out.relative_to(Path.cwd()) if out.is_relative_to(Path.cwd()) else out}")
        return

    for p in sorted(src.iterdir()):
        if p.suffix.lower() in (".docx", ".xlsx"):
            out = convert_one(p, dst)
            print(f"Wrote {out.relative_to(Path.cwd()) if out.is_relative_to(Path.cwd()) else out}")


if __name__ == "__main__":
    main()
